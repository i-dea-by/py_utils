"""
    понадобилось мне как-то посчитать в моих договорах общую сумму за год, бухалтерию-то не веду 🤪

    https://tokmakov.msk.ru/blog/item/78

"""
import pathlib

import docx

if __name__ == '__main__':
    pattern = '*.docx'
    path = pathlib.Path(r"t:\trashcan\-").glob(pattern)

    igogo = 0

    for item in path:
        # это если в каталоге есть отрытый файл и word создал временный файл
        if item.name[0] == '~':
            continue
        print(item, end=' ')
        doc = docx.Document(item)
        # в первой таблице в документе берем правую нижнюю ячейку
        table = doc.tables[0]
        summ = table.rows[-1].cells[-1].text
        if summ == '':
            # в некоторых файлах это была не первая таблица 🧐
            table = doc.tables[1]
            summ = table.rows[-1].cells[-1].text
        print('= ', summ)
        # разделители тысяч и копеек
        summ = summ.replace(',', '.').replace(' ', '')
        igogo += float(summ)
    print('Итого: ', igogo)
